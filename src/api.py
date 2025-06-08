from agents.agent import Agent 
from fastapi import FastAPI, File, Form, UploadFile, HTTPException, APIRouter, Body, Depends, BackgroundTasks, Query, Request
from fastapi.responses import FileResponse, HTMLResponse
from dotenv import load_dotenv
import threading
import os
from pydantic import BaseModel
import asyncio
import uvicorn
import time
from typing import List, Any, Callable, TypeVar, Awaitable, Optional, Dict, Union
import shutil
import concurrent.futures
import functools
from agents.umlAgent import UML_Agent
from fastapi.staticfiles import StaticFiles
from enum import Enum
from fastapi import Query
from agents.questionAgent import questionAgent, QuestionDifficulty, QuestionType
from agents.explainAgent import ExplainAgent
from agents.paperAgent import PaperAgent
from agents.testAgent import TestAgent, Language, TestType
from fastapi.middleware.cors import CORSMiddleware
import json
# 导入认证模块
from auth import auth_router, get_current_active_user, User
from utils.conversation_logger import ConversationLogger
from models.practice_history import PracticeHistory
from datetime import datetime
from models.review_plan import ReviewPlanManager
from agents.reviewplanAgent import ReviewPlanAgent



load_dotenv()
api_key = os.getenv("DASHSCOPE_API_KEY")
base_url = os.getenv("DASHSCOPE_BASE_URL")
model = 'qwen-plus'

# 确保必要的目录存在
KNOWLEDGE_DIR = os.path.join(os.getenv("PROJECT_PATH"), "knowledge_base")
VECTOR_STORE_DIR = os.path.join(os.getenv("PROJECT_PATH"), "VectorStore")

os.makedirs(KNOWLEDGE_DIR, exist_ok=True)
os.makedirs(VECTOR_STORE_DIR, exist_ok=True)

PROJECT_ROOT = os.getenv("PROJECT_PATH")

# 初始化对话记录器
PROJECT_PATH = os.getenv("PROJECT_PATH", os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
conversation_logger = ConversationLogger(PROJECT_PATH)

# 实例化习题历史记录管理器
practice_history = PracticeHistory(PROJECT_ROOT)

# 初始化复习计划管理器
review_plan_manager = ReviewPlanManager(PROJECT_ROOT)


# 全局变量
agent = Agent(api_key, base_url, model) 
agent_lock = threading.Lock()
agent_ready = threading.Event()
umlAgent = UML_Agent(api_key, base_url, model)
explainAgent = ExplainAgent(api_key, base_url, model)
question_agent = questionAgent(api_key, base_url, model)
paper_agent = PaperAgent()
test_agent = TestAgent(api_key, base_url, model)
review_plan_agent = ReviewPlanAgent(conversation_logger, practice_history, review_plan_manager)

agents = [agent, umlAgent, explainAgent, question_agent, paper_agent, test_agent, review_plan_agent]

type2agent = {
    "UmlAgent": umlAgent,
    "ExplainAgent": explainAgent,
    "QuestionAgent": question_agent,
    "PaperAgent": paper_agent,
    "TestAgent": test_agent,
}
# 保存全局事件循环引用
agent_loop = None

# 确保存放UML图片的目录存在
UML_STATIC_DIR = os.path.join(os.getenv("PROJECT_PATH"), "static")
os.makedirs(UML_STATIC_DIR, exist_ok=True)

# 确保文档输出目录存在
DOCS_DIR = os.path.join(os.getenv("PROJECT_PATH"), "static", "docs")
os.makedirs(DOCS_DIR, exist_ok=True)

class ChatRequest(BaseModel):
    message: str


# uml图类型
UML_TYPES = [
    "class", "sequence", "activity", "usecase", 
    "state", "component", "deployment", "object"
]

# 定义类型变量，用于泛型函数
T = TypeVar('T')

# 定义UML图类型枚举
class DiagramType(str, Enum):
    CLASS = "class"
    SEQUENCE = "sequence"
    ACTIVITY = "activity"
    USECASE = "usecase"
    STATE = "state"
    COMPONENT = "component"
    DEPLOYMENT = "deployment"
    OBJECT = "object"



# 定义解释风格枚举
    '''
    解释风格：
    - 严谨
    - 通俗
    - 专业
    - 简洁
    - 风趣
    '''
class ExplainStyle(str, Enum):
    STRICT = "STRICT"
    POPULAR = "POPULAR"
    PROFESSIONAL = "PROFESSIONAL"
    CONCISE = "CONCISE"
    FUNNY = "FUNNY"


# 工具函数
async def run_in_agent_thread(
    coro_func: Callable[..., Awaitable[T]], 
    *args, 
    timeout: int = 300,
    **kwargs
) -> T:
    """
    在Agent线程的事件循环中执行异步函数，并等待结果
    
    Args:
        coro_func: 要执行的异步函数
        *args: 传递给异步函数的位置参数
        timeout: 等待结果的超时时间（秒）
        **kwargs: 传递给异步函数的关键字参数
        
    Returns:
        异步函数的执行结果
        
    Raises:
        concurrent.futures.TimeoutError: 如果等待超时
        Exception: 如果异步函数执行出错
    """
    global agent_loop
    
    if agent_loop is None or agent_loop.is_closed():
        raise RuntimeError("Agent事件循环未创建或已关闭")
    
    # 创建Future对象，用于在线程间传递结果
    response_future = concurrent.futures.Future()
    
    def run_in_loop():
        """在agent事件循环中运行协程并设置Future结果"""
        try:
            # 创建协程
            coro = coro_func(*args, **kwargs)
            # 创建Task
            task = agent_loop.create_task(coro)
            
            # 添加回调函数以设置Future的结果
            def set_result(task):
                try:
                    result = task.result()
                    response_future.set_result(result)
                except Exception as e:
                    response_future.set_exception(e)
            
            task.add_done_callback(set_result)
        except Exception as e:
            response_future.set_exception(e)
    
    # 将函数调度到agent事件循环所在的线程中执行
    agent_loop.call_soon_threadsafe(run_in_loop)
    
    # 等待结果
    return await asyncio.get_event_loop().run_in_executor(
        None, 
        functools.partial(response_future.result, timeout=timeout)
    )

class ErrorTrackingRequest(BaseModel):
    question: str
    user_answer: str
    correct: bool

class StepByStepRequest(BaseModel):
    question: str

# 启动 Agent 的异步任务
async def start_agent():
    global agent
    try:
        for agt in agents:
            await agt.setup()
        agent_ready.set()
        print("Agent 初始化完成")
    except Exception as e:
        print(f"Agent 初始化失败: {e}")
        agent = None

# 在后台线程中启动 Agent
def background_start_agent():
    """在后台线程中初始化Agent并保持事件循环运行"""
    global  agent_loop
    
    # 创建新的事件循环
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    
    # 保存全局引用
    agent_loop = loop
    
    # 初始化Agent
    try:
        loop.run_until_complete(start_agent())
        print("Agent已初始化，事件循环继续运行")
        loop.run_forever()
    except Exception as e:
        print(f"Agent初始化或运行时出错: {e}")
    finally:
        # 仅在程序退出时关闭循环
        if loop.is_running():
            loop.stop()
        if not loop.is_closed():
            loop.close()

app = FastAPI()

# 添加CORS中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 在生产环境中，应该设置为具体的前端域名
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 挂载静态文件目录
app.mount("/static", StaticFiles(directory=UML_STATIC_DIR), name="static")

DIST_DIR = os.path.join(os.getenv("PROJECT_PATH"), "dist")

# 挂载dist目录
app.mount("/dist", StaticFiles(directory=DIST_DIR), name="dist")

# 集成认证路由
app.include_router(auth_router)


# @app.get("/")
# async def root():
#     """
#     API根路径，返回服务状态信息
    
#     Returns:
#         dict: 包含服务状态信息的字典
#     """
#     return {"message": "提供agent服务"}

@app.post("/chat")
async def chat(
    message: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    处理用户与Agent的对话请求
    
    Args:
        message (str): 用户发送的消息内容
        current_user (User): 当前登录的用户
        
    Returns:
        dict: 包含状态和回复消息的字典
        {
            "status": "success"/"error",
            "message": str
        }
    """
    if not agent or not agent_ready.is_set():
        return {"status": "error", "message": "Agent 尚未准备好，请稍后再试"}
    
    try:
        # 在Agent线程的事件循环中执行异步函数
        response = await run_in_agent_thread(agent.chat, message, timeout=300)
        
        # 记录对话
        agent_type = agent.__class__.__name__
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type=agent_type,
            query=message,
            response=response
        )
        
        return {"status": "success", "message": response['message']}
    except concurrent.futures.TimeoutError:
        return {"status": "error", "message": "请求超时，请尝试简化问题或稍后重试"}
    except Exception as e:
        print(f"处理消息时出错: {e}")
        return {"status": "error", "message": f"处理消息时出错: {str(e)}"}

@app.post("/create_or_update_index")
async def create_or_update_index(
    files: List[UploadFile] = File(...),
    name: str = Form(...),
    is_update: bool = Form(False),  # 添加此参数，表示是否为更新操作
    current_user: User = Depends(get_current_active_user)
):
    """
    创建或更新知识库索引
    
    Args:
        files (List[UploadFile]): 上传的文件列表
        name (str): 知识库名称
        is_update (bool): 是否为更新操作，True表示向现有知识库添加文件
        current_user (User): 当前登录的用户
        
    Returns:
        dict: 包含操作状态和消息的字典
        {
            "status": "success"/"error",
            "message": str
        }
    """
    if not name:
        raise HTTPException(status_code=400, detail="请提供知识库名称")

    # 检查知识库是否存在
    kb_dir = os.path.join(KNOWLEDGE_DIR, name)
    is_existing_kb = os.path.exists(kb_dir)
    
    # 如果是更新操作但知识库不存在，则返回错误
    if is_update and not is_existing_kb:
        return {"status": "error", "message": f"知识库 {name} 不存在，无法更新"}
    
    # 如果是创建操作但知识库已存在，则返回错误（除非明确指定为更新操作）
    if not is_update and is_existing_kb:
        return {"status": "error", "message": f"知识库 {name} 已存在，请使用更新功能或选择其他名称"}

    # 创建知识库目录（如果不存在）
    os.makedirs(kb_dir, exist_ok=True)
    
    # 保存上传的文件
    file_paths = []
    try:
        for file in files:
            print(f"处理文件: {file.filename}")
            dest_path = os.path.join(kb_dir, file.filename)
            
            # 检查文件是否已存在
            if os.path.exists(dest_path) and is_update:
                # 更新模式下，为重名文件添加时间戳
                file_name, file_ext = os.path.splitext(file.filename)
                timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
                new_filename = f"{file_name}_{timestamp}{file_ext}"
                dest_path = os.path.join(kb_dir, new_filename)
                print(f"文件已存在，重命名为: {new_filename}")
            
            # 读取上传的文件内容并写入目标位置
            with open(dest_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
                
            file_paths.append(dest_path)
            print(f"已保存文件到: {dest_path}")

    except Exception as e:
        return {"status": "error", "message": f"保存文件时出错: {str(e)}"}
    
    operation_type = "更新" if is_update else "创建"
    
    try:
        # 等待结果
        await run_in_agent_thread(agent.create_index, files_dir=kb_dir, label=name, timeout=120)
        return {"status": "success", "message": f"成功{operation_type}知识库: {name}，添加了 {len(file_paths)} 个文件"}
    except Exception as e:
        return {"status": "error", "message": f"{operation_type}向量存储时出错: {str(e)}"}

@app.get("/list_knowledge_bases")
async def list_knowledge_bases(current_user: User = Depends(get_current_active_user)):
    """
    获取所有知识库列表
    
    Args:
        current_user (User): 当前登录的用户
        
    Returns:
        dict: 包含知识库列表的字典
        {
            "status": "success",
            "knowledge_bases": List[str]
        }
    """
    return {"status": "success", "knowledge_bases": [os.path.basename(dir) for dir in os.listdir(KNOWLEDGE_DIR)]}

@app.post("/delete_knowledge_base")
async def delete_knowledge_base(
    name: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    删除指定的知识库
    
    Args:
        name (str): 要删除的知识库名称
        current_user (User): 当前登录的用户
        
    Returns:
        dict: 包含操作状态和消息的字典
        {
            "status": "success"/"error",
            "message": str
        }
    """
    if not name:
        raise HTTPException(status_code=400, detail="请提供知识库名称")
    
    try:
        await run_in_agent_thread(agent.delete_index, name, timeout=30)
        return {"status": "success", "message": f"成功删除知识库: {name}"}
    except Exception as e:
        return {"status": "error", "message": f"删除知识库时出错: {str(e)}"}


@app.post("/update_label")
async def update_label(
    name: str = Form(...),
    agent_type_list: List[str] = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    更新所有Agent使用的知识库标签
    
    Args:
        name (str): 知识库标签名称
        current_user (User): 当前登录的用户
        
    Returns:
        dict: 包含操作状态和消息的字典
        {
            "status": "success"/"error",
            "message": str
        }
    """
    if not name:
        raise HTTPException(status_code=400, detail="请提供知识库名称")
    
    if not agent or not agent_ready.is_set():
        return {"status": "error", "message": "Agent 尚未准备好，请稍后再试"}
    
    if agent_loop is None or agent_loop.is_closed():
        return {"status": "error", "message": "Agent事件循环未创建或已关闭"}


    # 检查label是否存在
    if name not in os.listdir(KNOWLEDGE_DIR):
        return {"status": "error", "message": f"知识库{name}不存在"}

    try:
        # 等待结果
        for agent_type in agent_type_list:
            agt = type2agent[agent_type]
            await run_in_agent_thread(agt.update_label, name, timeout=30)

        return {"status": "success", "message": f"成功更新知识库标签: {name}"}
    except Exception as e:
        return {"status": "error", "message": f"更新知识库标签时出错: {str(e)}"}

@app.post("/umlAgent/generate_uml")
async def generate_uml(
    query: str = Form(...),
    diagram_type: DiagramType = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    生成UML图
    
    Args:
        query (str): 用户的UML图生成请求
        diagram_type (DiagramType): UML图类型，支持的类型包括：class, sequence, activity等
        current_user (User): 当前登录的用户
        
    Returns:
        dict: 包含生成的UML图信息的字典
        {
            "status": "success"/"error",
            "message": str,
            "static_path": str,
        }
    """

    
    try:
        response = await run_in_agent_thread(
            umlAgent.generate_uml, 
            query=query, 
            diagram_type=diagram_type.value,
            timeout=300
        )
        
        # 记录对话
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="UmlAgent",
            query=f"生成{diagram_type.value}图：{query}",
            response=response
        )
        
        return {"status": "success", "message": response['message'], "static_path": f"http://localhost:8000/static/{diagram_type.value}/uml.png"}
    except Exception as e:
        print(f"生成UML图时出错: {e}")
        return {"status": "error", "message": f"生成UML图时出错: {str(e)}"}

@app.post("/explainAgent/explain")
async def explain(
    query: str = Form(...),
    style_label: ExplainStyle = Form(...),
    bing_search: bool = Form(False),
    current_user: User = Depends(get_current_active_user)
):
    """
    处理概念解释请求
    """
    try:
        # 使用位置参数调用
        response = await run_in_agent_thread(
            explainAgent.chat,
            query,                  # 第一个参数
            style_label.value,      # 第二个参数
            bing_search,            # 第四个参数
            timeout=120             # timeout是run_in_agent_thread的参数
        )



        # 记录对话
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="ExplainAgent",
            query=f"解释概念（{style_label.value}风格）：{query}",
            response=response
        )

        if response['status'] == 'success':
            return response
        else:
            return {
                "status": "error",
                "message": response['message']
            }
    except Exception as e:
        print(f"获取解释时出错: {e}")
        return {"status": "error", "message": f"获取解释时出错: {str(e)}"}

@app.post("/questionAgent/explain_question")
async def explain_question(
    question: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    解析软件工程习题
    """
    try:
        # 调用Agent解析题目
        response = await run_in_agent_thread(
            question_agent.explain_question,
            question
        )
        
        # 记录Agent响应
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="QuestionAgent",
            query=f"解释题目：{question}",
            response=response
        )

        return response
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/questionAgent/quick_answer")
async def quick_answer(
    question: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    快速回答软件工程相关问题
    """
    try:
        # 记录用户对话
        conversation_logger = ConversationLogger(current_user.username)
        conversation_id = await conversation_logger.log_conversation(
            agent_type="QuestionAgent",
            user_input=question,
            action="quick_answer"
        )
        
        # 调用Agent快速回答问题
        response = await run_in_agent_thread(
            question_agent.quick_answer,
            question
        )
        
        # 记录Agent响应
        await conversation_logger.update_conversation(
            conversation_id=conversation_id,
            agent_response=json.dumps(response)
        )
        
        return response
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/questionAgent/generate_practice_set")
async def generate_practice_set(
    topics: str = Form(...),
    num_questions: int = Form(5),
    difficulty: QuestionDifficulty = Form(QuestionDifficulty.MEDIUM),
    type: QuestionType = Form(QuestionType.MULTIPLE_CHOICE),
    current_user: User = Depends(get_current_active_user)
):
    """
    生成练习题

    data:
        {{
            "questions": [
                {{
                    "id": 1,
                    "type": "题目类型",
                    "question": "题目描述",
                    "options": ["选项1", "选项2", ...],  // 选择题必须提供
                    "reference_answer": "参考答案",
                    "analysis": "解题思路",
                    "topics": ["涉及知识点1", "涉及知识点2", ...]
                }},
                ...
            ],
            "total_points": 总分,
            "estimated_time": "预计完成时间（分钟）",
            "difficulty_distribution": {{
                "easy": 简单题数量,
                "medium": 中等题数量,
                "hard": 困难题数量
            }}
        }}
    """
    try:
        # 将接收到的字符串转换为列表
        topic_list = [topic.strip() for topic in topics.split(",") if topic.strip()]
        
        # 调用question_agent时传入处理后的列表
        result = await run_in_agent_thread(
            question_agent.generate_practice_set,
            topic_list,
            num_questions,
            difficulty,
            type
        )

        # 记录对话
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="QuestionAgent",
            query=f"生成练习集：主题={topics}，数量={num_questions}，难度={difficulty.value}",
            response=result
        )
        return result
    except Exception as e:
        print(f"生成练习集时出错: {e}")
        return {"status": "error", "message": f"生成练习集时出错: {str(e)}"}

@app.post("/questionAgent/grade_practice_set")
async def grade_practice_set(
    practice_set: str = Form(...),  # 题目集的JSON字符串
    student_answers: str = Form(...),  # 学生答案集的JSON字符串
    reference_answers: str = Form(...),  # 参考答案集的JSON字符串
    current_user: User = Depends(get_current_active_user)
):
    """
    批改练习题集
    
    Args:
        practice_set: 题目集的JSON字符串
        reference_answers: 参考答案集的JSON字符串
        
    Returns:
        dict: 包含批改结果的字典
        {
            "status": "success"/"error",
            "data": {
                "score": float,  # 总分
                "scoring_points": [  # 得分点详情
                    {
                        "id": str,  # 题目ID
                        "point": str,  # 得分/失分点描述
                        "score": float,  # 得分
                        "deduction": float  # 扣分（如果是失分点）
                    },
                    ...
                ],
                "comments": str,  # 总体评价
                "suggestions": List[str],  # 改进建议
                "highlights": List[str]  # 亮点
            }
        }
    """
    try:
        # 解析JSON字符串
        try:
            practice_set_data = json.loads(practice_set)
            reference_answers_data = json.loads(reference_answers)
            student_answers_data = json.loads(student_answers)
        except json.JSONDecodeError:
            return {
                "status": "error",
                "message": "输入数据格式错误，请确保是有效的JSON格式"
            }
        
        # 调用question_agent的批改方法
        result = await question_agent.grade_practice_set(
            practice_set=practice_set_data,
            student_answers=student_answers_data,
            reference_answers=reference_answers_data    
        )
        
        if result["status"] == "success":
            return {
                "status": "success",
                "data": result["message"]
            }
        else:
            return {
                "status": "error",
                "message": result["message"]
            }
            
    except Exception as e:
        return {
            "status": "error",
            "message": f"批改练习题集时出错: {str(e)}"
        }

@app.post("/paperAgent/search_papers")
async def search_papers(
    topic: str = Form(...), 
    max_results: int = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    搜索论文
    
    Args:
        topic (str): 搜索主题
        max_results (int): 最大结果数
        current_user (User): 当前登录的用户
        
    Returns:
        dict: 包含搜索结果的字典
    """
    if not paper_agent or not agent_ready.is_set():
        return {"status": "error", "message": "Paper Agent 尚未准备好，请稍后再试"}
    
    try:
        response = await run_in_agent_thread(
            paper_agent.search_papers_by_topic,
            topic=topic,
            max_results=max_results,
            timeout=300
        )
        
        # 记录对话
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="PaperAgent",
            query=f"搜索论文：{topic}，最大结果数={max_results}",
            response=response
        )
        
        return response
    except Exception as e:
        print(f"搜索论文时出错: {e}")
        return {"status": "error", "message": f"搜索论文时出错: {str(e)}"}

@app.post("/paperAgent/download_and_read_paper")
async def download_and_read_paper(
    paper_id: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    获取论文详情
    """
    try:
        result = await run_in_agent_thread(paper_agent.download_and_read_paper, paper_id, timeout=120)
        return {"status": "success", "message": result['message']}
    except Exception as e:
        return {"status": "error", "message": f"下载和阅读论文时出错: {str(e)}"}

@app.post("/paperAgent/list_and_organize_papers")
async def list_and_organize_papers(
    current_user: User = Depends(get_current_active_user)
):
    """
    列出并组织论文
    """
    try:
        result = await run_in_agent_thread(paper_agent.list_and_organize_papers, timeout=120)
        return {"status": "success", "message": result['message']}
    except Exception as e:
        return {"status": "error", "message": f"列出和组织论文时出错: {str(e)}"}        
    

@app.post("/paperAgent/analyze_paper_for_project")
async def analyze_paper_for_project(
    paper_id: str = Form(...), 
    project_description: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    分析论文对特定项目的应用价值
    """
    try:
        result = await run_in_agent_thread(paper_agent.analyze_paper_for_project, paper_id, project_description, timeout=120)
        return {"status": "success", "message": result['message']}
    except Exception as e:
        return {"status": "error", "message": f"分析论文对特定项目的应用价值时出错: {str(e)}"}
     

@app.post("/paperAgent/recommend_learning_path")
async def recommend_learning_path(
    topic: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    推荐学习路径
    """
    try:
        result = await run_in_agent_thread(paper_agent.recommend_learning_path, topic, timeout=120)
        return {"status": "success", "message": result['message']}
    except Exception as e:
        return {"status": "error", "message": f"推荐学习路径时出错: {str(e)}"}



# 添加新的下载端点
@app.get("/download/{filename}")
async def download_file(
    filename: str,
    current_user: User = Depends(get_current_active_user)
):
    """
    下载生成的解释文件
    
    Args:
        filename: 文件名
        
    Returns:
        FileResponse: 文件下载响应
    """
    file_path = os.path.join(DOCS_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
        
    return FileResponse(
        file_path,
        media_type='application/octet-stream',
        filename=filename
    )

@app.post("/conversations")
async def get_user_conversations(
    limit: Optional[int] = Query(None, description="返回的最大记录数"),
    agent_type: Optional[str] = Query(None, description="可选的Agent类型过滤条件"),
    current_user: User = Depends(get_current_active_user)
):
    """
    获取用户的对话历史记录
    
    Args:
        limit: 返回的最大记录数，默认20条
        agent_type: 可选的Agent类型过滤条件
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含用户对话记录的字典
    """
    try:
        conversations = conversation_logger.get_user_conversations(
            user_id=current_user.id,
            limit=limit,
            agent_type=agent_type
        )
        
        return {
            "status": "success",
            "conversations": conversations,
            "total": len(conversations)
        }
    except Exception as e:
        print(f"获取用户对话记录时出错: {e}")
        return {"status": "error", "message": f"获取用户对话记录时出错: {str(e)}"}

@app.delete("/conversations/{conversation_id}")
async def delete_conversation(
    conversation_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """
    删除特定的对话记录
    
    Args:
        conversation_id: 要删除的对话ID
        current_user: 当前登录的用户
        
    Returns:
        dict: 操作结果
    """
    try:
        # 验证conversation_id是否属于当前用户
        user_id_str = conversation_id.split('_')[0]
        try:
            conversation_user_id = int(user_id_str)
            if conversation_user_id != current_user.id:
                return {"status": "error", "message": "无权删除此对话记录"}
        except ValueError:
            return {"status": "error", "message": "无效的对话ID"}
            
        # 获取用户的所有对话记录
        conversations = conversation_logger._load_user_conversations(current_user.id)
        
        # 找到并删除指定ID的对话
        before_count = len(conversations["conversations"])
        conversations["conversations"] = [
            conv for conv in conversations["conversations"] 
            if conv["id"] != conversation_id
        ]
        after_count = len(conversations["conversations"])
        
        if before_count == after_count:
            return {"status": "error", "message": "未找到指定的对话记录"}
            
        # 保存更新后的对话记录
        conversation_logger._save_user_conversations(current_user.id, conversations)
        
        return {
            "status": "success", 
            "message": "已成功删除对话记录"
        }
    except Exception as e:
        print(f"删除对话记录时出错: {e}")
        return {"status": "error", "message": f"删除对话记录时出错: {str(e)}"}

@app.delete("/conversations")
async def delete_all_conversations(
    agent_type: Optional[str] = Query(None, description="可选的Agent类型过滤条件"),
    current_user: User = Depends(get_current_active_user)
):
    """
    删除用户的所有对话记录
    
    Args:
        agent_type: 可选的Agent类型过滤条件，如果提供则只删除该类型的对话
        current_user: 当前登录的用户
        
    Returns:
        dict: 操作结果
    """
    try:
        # 获取用户的所有对话记录
        conversations = conversation_logger._load_user_conversations(current_user.id)
        
        if agent_type:
            # 只删除特定类型的对话
            before_count = len(conversations["conversations"])
            conversations["conversations"] = [
                conv for conv in conversations["conversations"] 
                if conv["agent_type"] != agent_type
            ]
            after_count = len(conversations["conversations"])
            deleted_count = before_count - after_count
            
            # 保存更新后的对话记录
            conversation_logger._save_user_conversations(current_user.id, conversations)
            
            return {
                "status": "success", 
                "message": f"已成功删除 {deleted_count} 条 {agent_type} 类型的对话记录"
            }
        else:
            # 删除所有对话
            deleted_count = len(conversations["conversations"])
            conversations["conversations"] = []
            
            # 保存更新后的对话记录
            conversation_logger._save_user_conversations(current_user.id, conversations)
            
            return {
                "status": "success", 
                "message": f"已成功删除所有 {deleted_count} 条对话记录"
            }
    except Exception as e:
        print(f"删除对话记录时出错: {e}")
        return {"status": "error", "message": f"删除对话记录时出错: {str(e)}"}

@app.post("/testAgent/generate_test_cases")
async def generate_test_cases(
    code: str = Form(...),
    language: Language = Form(...),
    test_type: TestType = Form(...),
    description: str = Form(""),
    current_user: User = Depends(get_current_active_user)
):
    """
    为给定代码生成测试用例
    
    Args:
        code: 要测试的源代码
        language: 编程语言
        test_type: 测试类型
        description: 代码功能描述(可选)
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含测试用例代码和解释的字典
    """
    if not test_agent or not agent_ready.is_set():
        return {"status": "error", "message": "Test Agent 尚未准备好，请稍后再试"}
    
    try:
        response = await run_in_agent_thread(
            test_agent.generate_test_cases,
            code=code,
            language=language,
            test_type=test_type,
            description=description,
            timeout=300
        )
        
        # 记录对话
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="TestAgent",
            query=f"生成{language.value}代码的{test_type.value}测试用例",
            response=response
        )
        
        return response
    except Exception as e:
        print(f"生成测试用例时出错: {e}")
        return {"status": "error", "message": f"生成测试用例时出错: {str(e)}"}


@app.post("/testAgent/analyze_code_for_testability")
async def analyze_code_for_testability(
    code: str = Form(...),
    language: Language = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    分析代码的可测试性并提供改进建议
    
    Args:
        code: 要分析的代码
        language: 编程语言
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含可测试性分析结果的字典
    """
    if not test_agent or not agent_ready.is_set():
        return {"status": "error", "message": "Test Agent 尚未准备好，请稍后再试"}
    
    try:
        response = await run_in_agent_thread(
            test_agent.analyze_code_for_testability,
            code=code,
            language=language,
            timeout=300
        )
        
        # 记录对话
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="TestAgent",
            query=f"分析{language.value}代码的可测试性",
            response=response
        )
        
        return response
    except Exception as e:
        print(f"分析代码可测试性时出错: {e}")
        return {"status": "error", "message": f"分析代码可测试性时出错: {str(e)}"}


@app.post("/testAgent/evaluate_test_coverage")
async def evaluate_test_coverage(
    code: str = Form(...),
    tests: str = Form(...),
    language: Language = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    评估测试用例对代码的覆盖程度
    
    Args:
        code: 源代码
        tests: 测试代码
        language: 编程语言
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含测试覆盖率评估的字典
    """
    if not test_agent or not agent_ready.is_set():
        return {"status": "error", "message": "Test Agent 尚未准备好，请稍后再试"}
    
    try:
        response = await run_in_agent_thread(
            test_agent.evaluate_test_coverage,
            code=code,
            tests=tests,
            language=language,
            timeout=300
        )
        
        # 记录对话
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="TestAgent",
            query=f"评估{language.value}测试代码的覆盖率",
            response=response
        )
        
        return response
    except Exception as e:
        print(f"评估测试覆盖率时出错: {e}")
        return {"status": "error", "message": f"评估测试覆盖率时出错: {str(e)}"}


@app.post("/questionAgent/download_practice_set")
async def download_practice_set(
    practice_set: str = Form(...),  # 题目集的JSON字符串
    format: str = Form("pdf"),  # 下载格式，支持pdf或docx
    current_user: User = Depends(get_current_active_user)
):
    """
    下载习题集
    
    Args:
        practice_set: 题目集的JSON字符串
        format: 下载格式，支持pdf或docx
        current_user: 当前登录的用户
        
    Returns:
        FileResponse: 文件下载响应
    """
    try:
        # 解析JSON字符串
        try:
            practice_set_data = json.loads(practice_set)
        except json.JSONDecodeError:
            return {
                "status": "error",
                "message": "输入数据格式错误，请确保是有效的JSON格式"
            }
        
        # 生成唯一的文件名
        timestamp = int(time.time())
        filename = f"practice_set_{current_user.id}_{timestamp}"
        
        # 文件路径
        file_path = os.path.join(DOCS_DIR, f"{filename}.{format}")
        
        # 根据格式生成文件
        if format == "pdf":
            await generate_pdf(practice_set_data, file_path)
        elif format == "docx":
            await generate_docx(practice_set_data, file_path)
        else:
            return {"status": "error", "message": "不支持的格式，只支持pdf或docx"}
        
        # 返回文件下载响应
        return FileResponse(
            file_path,
            media_type='application/octet-stream',
            filename=f"练习题集_{timestamp}.{format}"
        )
    except Exception as e:
        return {"status": "error", "message": f"下载习题集时出错: {str(e)}"}

async def generate_pdf(practice_set_data, file_path):
    """
    生成PDF文件
    
    Args:
        practice_set_data: 习题集数据
        file_path: 输出文件路径
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import inch
    
    # 注册中文字体
    try:
        # 尝试注册思源黑体（Source Han Sans）
        pdfmetrics.registerFont(TTFont('SourceHanSans', 'SourceHanSans-Regular.ttf'))
        font_name = 'SourceHanSans'
    except:
        try:
            # 尝试注册宋体
            pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
            font_name = 'SimSun'
        except:
            # 如果都失败，使用默认字体
            font_name = 'Helvetica'
            print("警告：未能加载中文字体，文档可能无法正确显示中文")
    
    # 创建自定义样式
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ChineseTitle',
        fontName=font_name,
        fontSize=18,
        leading=22,
        alignment=1  # 居中
    ))
    styles.add(ParagraphStyle(
        name='ChineseHeading',
        fontName=font_name,
        fontSize=14,
        leading=18
    ))
    styles.add(ParagraphStyle(
        name='ChineseNormal',
        fontName=font_name,
        fontSize=10,
        leading=14
    ))
    
    # 创建PDF文档
    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # 创建内容列表
    content = []
    
    # 添加标题
    title = Paragraph("软件工程练习题集", styles['ChineseTitle'])
    content.append(title)
    content.append(Spacer(1, 12))
    
    # 添加题目
    for i, question in enumerate(practice_set_data.get("questions", [])):
        # 题目标题
        question_title = Paragraph(f"{i+1}. {question.get('question', '未提供题目')}", styles['ChineseHeading'])
        content.append(question_title)
        content.append(Spacer(1, 6))
        
        # 选项（如果有）
        if question.get("options"):
            for j, option in enumerate(question["options"]):
                option_text = Paragraph(f"{option}", styles['ChineseNormal'])
                content.append(option_text)
            content.append(Spacer(1, 6))
        
        # 参考答案
        answer_title = Paragraph("参考答案:", styles['ChineseHeading'])
        content.append(answer_title)
        answer_text = Paragraph(question.get("reference_answer", "未提供答案"), styles['ChineseNormal'])
        content.append(answer_text)
        content.append(Spacer(1, 6))
        
        # 解析
        if question.get("analysis"):
            analysis_title = Paragraph("解析:", styles['ChineseHeading'])
            content.append(analysis_title)
            analysis_text = Paragraph(question.get("analysis", ""), styles['ChineseNormal'])
            content.append(analysis_text)
        
        content.append(Spacer(1, 12))
    
    # 构建PDF
    doc.build(content)

async def generate_docx(practice_set_data, file_path):
    """
    生成DOCX文件
    
    Args:
        practice_set_data: 习题集数据
        file_path: 输出文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 创建文档
        doc = Document()
        
        # 设置中文字体
        chinese_font = '宋体'  # 或者使用'微软雅黑', '黑体'等
        
        # 添加标题
        title = doc.add_heading("软件工程练习题集", 0)
        for run in title.runs:
            run.font.name = chinese_font
        
        # 添加题目
        for i, question in enumerate(practice_set_data.get("questions", [])):
            # 题目标题
            heading = doc.add_heading(f"{i+1}. {question.get('question', '未提供题目')}", 2)
            for run in heading.runs:
                run.font.name = chinese_font
            
            # 选项（如果有）
            if question.get("options"):
                for option in question["options"]:
                    para = doc.add_paragraph(option)
                    for run in para.runs:
                        run.font.name = chinese_font
            
            # 参考答案
            answer_para = doc.add_paragraph()
            answer_run = answer_para.add_run("参考答案:")
            answer_run.bold = True
            answer_run.font.name = chinese_font
            
            answer_text = doc.add_paragraph(question.get("reference_answer", "未提供答案"))
            for run in answer_text.runs:
                run.font.name = chinese_font
            
            # 解析
            if question.get("analysis"):
                analysis_para = doc.add_paragraph()
                analysis_run = analysis_para.add_run("解析:")
                analysis_run.bold = True
                analysis_run.font.name = chinese_font
                
                analysis_text = doc.add_paragraph(question.get("analysis", ""))
                for run in analysis_text.runs:
                    run.font.name = chinese_font
            
            # 添加分隔线
            doc.add_paragraph("-----------------------------------")
        
        # 保存文档
        doc.save(file_path)
    except ImportError:
        # 如果python-docx不可用，使用简单的文本文件
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("软件工程练习题集\n\n")
            
            for i, question in enumerate(practice_set_data.get("questions", [])):
                f.write(f"{i+1}. {question.get('question', '未提供题目')}\n\n")
                
                if question.get("options"):
                    for j, option in enumerate(question["options"]):
                        f.write(f"{option}\n")
                    f.write("\n")
                
                f.write(f"参考答案: {question.get('reference_answer', '未提供答案')}\n\n")
                
                if question.get("analysis"):
                    f.write(f"解析: {question.get('analysis', '')}\n\n")
                
                f.write("-----------------------------------\n\n")

@app.get("/questionAgent/practice_history")
async def get_practice_history(
    limit: Optional[int] = Query(20, description="返回的最大记录数"),
    current_user: User = Depends(get_current_active_user)
):
    """
    获取用户的习题历史记录
    
    Args:
        limit: 返回的最大记录数，默认20条
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含用户习题历史记录的字典
    """
    try:
        history = practice_history.get_user_history(current_user.id, limit)
        
        return {
            "status": "success",
            "history": history,
            "total": len(history)
        }
    except Exception as e:
        print(f"获取用户习题历史记录时出错: {e}")
        return {"status": "error", "message": f"获取用户习题历史记录时出错: {str(e)}"}

@app.post("/questionAgent/save_practice_history")
async def save_practice_history(
    topics: str = Form(...),
    count: int = Form(...),
    difficulty: str = Form(...),
    type: str = Form(...),
    questions: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    保存习题历史记录
    
    Args:
        topics: 知识点列表，逗号分隔
        count: 题目数量
        difficulty: 难度
        type: 题目类型
        questions: 习题JSON字符串
        current_user: 当前登录的用户
        
    Returns:
        dict: 操作结果
    """
    try:
        # 解析JSON字符串
        try:
            questions_data = json.loads(questions)
        except json.JSONDecodeError:
            return {
                "status": "error",
                "message": "输入数据格式错误，请确保是有效的JSON格式"
            }
        
        # 创建历史记录项
        item_id = f"{current_user.id}_{int(time.time())}"
        topics_list = [topic.strip() for topic in topics.split(",") if topic.strip()]
        
        history_item = {
            "id": item_id,
            "user_id": current_user.id,
            "date": datetime.now().isoformat(),
            "topics": topics_list,
            "count": count,
            "difficulty": difficulty,
            "type": type,
            "questions": questions_data
        }
        
        # 添加到历史记录
        success = practice_history.add_history_item(current_user.id, history_item)
        
        if success:
            return {"status": "success", "message": "习题历史记录保存成功", "id": item_id}
        else:
            return {"status": "error", "message": "习题历史记录保存失败"}
    except Exception as e:
        print(f"保存习题历史记录时出错: {e}")
        return {"status": "error", "message": f"保存习题历史记录时出错: {str(e)}"}

@app.delete("/questionAgent/practice_history/{item_id}")
async def delete_practice_history(
    item_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """
    删除习题历史记录
    
    Args:
        item_id: 历史记录项ID
        current_user: 当前登录的用户
        
    Returns:
        dict: 操作结果
    """
    try:
        # 验证item_id是否属于当前用户
        user_id_str = item_id.split('_')[0]
        try:
            item_user_id = int(user_id_str)
            if item_user_id != current_user.id:
                return {"status": "error", "message": "无权删除此历史记录"}
        except ValueError:
            return {"status": "error", "message": "无效的历史记录ID"}
            
        # 删除历史记录
        success = practice_history.delete_history_item(current_user.id, item_id)
        
        if success:
            return {"status": "success", "message": "历史记录删除成功"}
        else:
            return {"status": "error", "message": "历史记录删除失败，可能记录不存在"}
    except Exception as e:
        print(f"删除习题历史记录时出错: {e}")
        return {"status": "error", "message": f"删除习题历史记录时出错: {str(e)}"}

@app.delete("/questionAgent/practice_history")
async def clear_practice_history(
    current_user: User = Depends(get_current_active_user)
):
    """
    清空用户的习题历史记录
    
    Args:
        current_user: 当前登录的用户
        
    Returns:
        dict: 操作结果
    """
    try:
        success = practice_history.clear_user_history(current_user.id)
        
        if success:
            return {"status": "success", "message": "历史记录清空成功"}
        else:
            return {"status": "error", "message": "历史记录清空失败"}
    except Exception as e:
        print(f"清空习题历史记录时出错: {e}")
        return {"status": "error", "message": f"清空习题历史记录时出错: {str(e)}"}

@app.get("/questionAgent/practice_history/{item_id}")
async def get_practice_history_item(
    item_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """
    获取特定的习题历史记录
    
    Args:
        item_id: 历史记录项ID
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含历史记录项的字典
    """
    try:
        # 验证item_id是否属于当前用户
        user_id_str = item_id.split('_')[0]
        try:
            item_user_id = int(user_id_str)
            if item_user_id != current_user.id:
                return {"status": "error", "message": "无权查看此历史记录"}
        except ValueError:
            return {"status": "error", "message": "无效的历史记录ID"}
            
        # 获取所有历史记录
        history = practice_history.get_user_history(current_user.id)
        
        # 查找特定记录
        item = next((item for item in history if item.get('id') == item_id), None)
        
        if item:
            return {"status": "success", "data": item}
        else:
            return {"status": "error", "message": "未找到指定的历史记录"}
    except Exception as e:
        print(f"获取习题历史记录项时出错: {e}")
        return {"status": "error", "message": f"获取习题历史记录项时出错: {str(e)}"}

@app.post("/reviewAgent/generate_plan")
async def generate_review_plan(
    current_user: User = Depends(get_current_active_user)
):
    """
    生成复习计划
    
    基于用户的历史对话记录、学习笔记和错题集，生成个性化复习计划
    
    Args:
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含复习计划的字典
    """
    try:
        response = await run_in_agent_thread(
            review_plan_agent.generate_review_plan,
            user_id=current_user.id,
            username=current_user.username,
            timeout=300
        )
        
        # 记录对话
        conversation_logger.log_conversation(
            user_id=current_user.id,
            username=current_user.username,
            agent_type="ReviewPlanAgent",
            query="生成复习计划",
            response=response
        )
        
        return response
    except Exception as e:
        print(f"生成复习计划时出错: {e}")
        return {"status": "error", "message": f"生成复习计划时出错: {str(e)}"}

@app.get("/reviewAgent/plans")
async def get_user_plans(
    limit: Optional[int] = Query(None, description="返回的最大记录数"),
    current_user: User = Depends(get_current_active_user)
):
    """
    获取用户的复习计划列表
    
    Args:
        limit: 返回的最大记录数
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含用户复习计划列表的字典
    """
    try:
        plans = review_plan_manager.get_user_plans(current_user.id, limit)
        
        return {
            "status": "success",
            "plans": plans,
            "total": len(plans)
        }
    except Exception as e:
        print(f"获取用户复习计划列表时出错: {e}")
        return {"status": "error", "message": f"获取用户复习计划列表时出错: {str(e)}"}

@app.get("/reviewAgent/plans/{plan_id}")
async def get_plan_detail(
    plan_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """
    获取复习计划详情
    
    Args:
        plan_id: 计划ID
        current_user: 当前登录的用户
        
    Returns:
        dict: 包含复习计划详情的字典
    """
    try:
        plan = review_plan_manager.get_plan_by_id(current_user.id, plan_id)
        
        if not plan:
            return {"status": "error", "message": "未找到指定的复习计划"}
        
        return {
            "status": "success",
            "plan": plan
        }
    except Exception as e:
        print(f"获取复习计划详情时出错: {e}")
        return {"status": "error", "message": f"获取复习计划详情时出错: {str(e)}"}

@app.put("/reviewAgent/plans/{plan_id}/steps/{step_id}")
async def update_step_status(
    plan_id: str,
    step_id: str,
    is_completed: bool = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    """
    更新复习计划步骤状态
    
    Args:
        plan_id: 计划ID
        step_id: 步骤ID
        is_completed: 是否已完成
        current_user: 当前登录的用户
        
    Returns:
        dict: 操作结果
    """
    try:
        success = review_plan_manager.update_step_status(
            current_user.id, 
            plan_id, 
            step_id, 
            is_completed
        )
        
        if success:
            return {"status": "success", "message": "步骤状态更新成功"}
        else:
            return {"status": "error", "message": "步骤状态更新失败，可能计划或步骤不存在"}
    except Exception as e:
        print(f"更新复习计划步骤状态时出错: {e}")
        return {"status": "error", "message": f"更新复习计划步骤状态时出错: {str(e)}"}

@app.delete("/reviewAgent/plans/{plan_id}")
async def delete_plan(
    plan_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """
    删除复习计划
    
    Args:
        plan_id: 计划ID
        current_user: 当前登录的用户
        
    Returns:
        dict: 操作结果
    """
    try:
        success = review_plan_manager.delete_plan(current_user.id, plan_id)
        
        if success:
            return {"status": "success", "message": "复习计划删除成功"}
        else:
            return {"status": "error", "message": "复习计划删除失败，可能计划不存在"}
    except Exception as e:
        print(f"删除复习计划时出错: {e}")
        return {"status": "error", "message": f"删除复习计划时出错: {str(e)}"}


# 后端 API 添加 /api 前缀
app_with_prefix = FastAPI()
app_with_prefix.mount("/api", app)

# 挂载静态资源
app_with_prefix.mount("/assets", StaticFiles(directory=os.path.join(DIST_DIR, "assets")), name="assets")

# 处理根路径，返回 index.html
@app_with_prefix.get("/", response_class=HTMLResponse)
async def read_root():
    try:
        index_path = os.path.join(DIST_DIR, "index.html")
        print(f"尝试读取首页: {index_path}")
        with open(index_path, mode="r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        print(f"读取首页出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")

# 处理所有其他路径，支持前端路由
@app_with_prefix.get("/{full_path:path}", response_class=HTMLResponse)
async def serve_spa(request: Request, full_path: str):
    print(f"处理请求路径: {full_path}")
    
    # 检查是否是直接静态文件请求（如 favicon.ico）
    if "." in full_path.split("/")[-1]:
        file_path = os.path.join(DIST_DIR, full_path)
        print(f"检查静态文件: {file_path}")
        if os.path.exists(file_path) and os.path.isfile(file_path):
            print(f"提供静态文件: {file_path}")
            return FileResponse(file_path)
        else:
            print(f"静态文件不存在: {file_path}")
    
    # 所有其他请求返回 index.html（支持前端路由）
    try:
        index_path = os.path.join(DIST_DIR, "index.html")
        print(f"返回SPA首页: {index_path}")
        with open(index_path, mode="r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        print(f"读取SPA首页出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")

# 使用新的应用实例
app = app_with_prefix

if __name__ == "__main__":
    # 启动后台线程
    agent_thread = threading.Thread(target=background_start_agent, daemon=True)
    agent_thread.start()
    
    # 等待agent初始化
    timeout = 120  
    start_time = time.time()
    while not agent_ready.is_set() and time.time() - start_time < timeout:
        time.sleep(0.5)
    
    if not agent_ready.is_set():
        print("警告: Agent初始化超时，API服务可能无法正常工作")
    
    # 禁用uvicorn的热重载功能
    uvicorn.run(app, host="127.0.0.1", port=8000, reload=False)