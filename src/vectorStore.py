from llama_index.core import VectorStoreIndex,Settings,SimpleDirectoryReader,load_index_from_storage
from llama_index.embeddings.dashscope import (
    DashScopeEmbedding,
    DashScopeTextEmbeddingModels,
    DashScopeTextEmbeddingType,
)
from llama_index.core.schema import TextNode
import os 
from dotenv import load_dotenv
from llama_index.core.storage import StorageContext
import shutil
# from utils.splitter import split_questions
from llama_index.core import Document

load_dotenv()

DB_PATH = os.getenv("PROJECT_PATH") + "/VectorStore"
KB_PATH = os.getenv("PROJECT_PATH") + "/knowledge_base"

class VectorStore:
    def __init__(self, index_path: str = DB_PATH, model_name: str = DashScopeTextEmbeddingModels.TEXT_EMBEDDING_V2, type: str = DashScopeTextEmbeddingType.TEXT_TYPE_DOCUMENT):
        
            

        self.index_path = index_path
        os.makedirs(self.index_path, exist_ok=True)

        self.embedding_model = DashScopeEmbedding(
        api_key=os.getenv("DASHSCOPE_API_KEY"),
        model_name=model_name,
        text_type=type,
        )
        Settings.embed_model = self.embedding_model

    def question_create_index(self, file_path: str, label: str):
        # 确认路径存在
        if not os.path.exists(file_path):
            raise ValueError(f"文件路径不存在: {file_path}")

        # 读取所有文件内容
        all_text = ""
        if os.path.isdir(file_path):
            for fname in os.listdir(file_path):
                if fname.endswith(".txt"):
                    with open(os.path.join(file_path, fname), "r", encoding="utf-8") as f:
                        all_text += f.read() + "\n"
                elif fname.endswith(".docx"):
                    try:
                        import docx
                        doc = docx.Document(os.path.join(file_path, fname))
                        for para in doc.paragraphs:
                            all_text += para.text + "\n"
                    except ImportError:
                        print("请安装 python-docx 包以支持 .docx 文件")
                        raise
        else:
            if file_path.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    all_text = f.read()
            elif file_path.endswith(".docx"):
                try:
                    import docx
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        all_text += para.text + "\n"
                except ImportError:
                    print("请安装 python-docx 包以支持 .docx 文件")
                    raise

        # 按题目分割

        question_chunks = split_questions(all_text)
        print(f"分块数量: {len(question_chunks)}")
        for i, chunk in enumerate(question_chunks[:3]):
            print(f"第{i+1}块内容预览:\n{chunk[:200]}\n")
        # # ... existing code ...
        # question_chunks = split_questions(all_text)
        # # 构造 TextNode 列表
        nodes = [Document(text=chunk) for chunk in question_chunks if chunk.strip()]

        # 构建索引
        index = VectorStoreIndex.from_documents(
            documents=nodes,
            embedding=self.embedding_model,
        )
        db_path = os.path.join(self.index_path, label)
        index.storage_context.persist(db_path)
        print(f"向量数据库创建成功，路径为{db_path}")
    
    def load_index(self, label: str):

        db_path = os.path.join(self.index_path, label)
        print(f"向量数据库路径为：{db_path}")
        if not os.path.exists(db_path):
            raise ValueError(f"向量数据库路径不存在: {db_path}")
        index = load_index_from_storage(
            storage_context=StorageContext.from_defaults(persist_dir=db_path),
        )
        return index
    

    def delete_index(self, label: str):
        kb_dir = os.path.join(KB_PATH, label)
        vs_dir = os.path.join(self.index_path, label)

        if os.path.exists(kb_dir):
            shutil.rmtree(kb_dir)
        
        else:
            raise ValueError(f"知识库不存在: {label}")

        if os.path.exists(vs_dir):
            shutil.rmtree(vs_dir)
        
        else:
            raise ValueError(f"向量数据库路径不存在: {vs_dir}") 
        
        return 1

    def list_label(self):
        return os.listdir(self.index_path)

if __name__ == "__main__":
    vector_store = VectorStore()
    vector_store.create_index(file_path=f"{os.getcwd()}/src/knowledge_base", label="knowledge_base")